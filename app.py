import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import json
import os
import sys
import subprocess
import webbrowser
import PyPDF2
from docx import Document
from io import BytesIO
import streamlit.components.v1 as components


# -------------------------------------------------
# 🔒 LOGIN-SCHUTZ
# -------------------------------------------------
VALID_USERS = {
    "jonathan": "IchBinJon",
    "Anna-Lena": "IchBinAnn",
    "lara": "IchBinLara",
}

def login_page():
    st.title("🔐 Login – Uni-Dashboard")
    username = st.text_input("Benutzername")
    password = st.text_input("Passwort", type="password")

    if st.button("Einloggen"):
        if username in VALID_USERS and VALID_USERS[username] == password:
            st.session_state["logged_in"] = True
            st.session_state["user"] = username
            st.success("Erfolgreich eingeloggt! 🎉")
            st.rerun()
        else:
            st.error("❌ Benutzername oder Passwort falsch")


if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if not st.session_state["logged_in"]:
    login_page()
    st.stop()


# -------------------------------------------------
# 👋 Begrüßungsbanner mit Tageszeit
# -------------------------------------------------
user = st.session_state.get("user", "Unbekannt")
welcome_name = user.capitalize()

hour = datetime.now().hour
if hour < 11:
    greeting = "🌅 Guten Morgen"
elif hour < 17:
    greeting = "☀️ Guten Tag"
else:
    greeting = "🌙 Guten Abend"

st.markdown(
    f"""
    <div style='
        background-color:#f0f2f6;
        padding:18px;
        border-radius:12px;
        margin-bottom:20px;
        border-left: 6px solid #4a90e2;
        font-size:20px;
    '>
        <b>{greeting}, {welcome_name} – Willkommen zurück! 👋</b>
    </div>
    """,
    unsafe_allow_html=True
)

# -------------------------------------------------
# Nutzerabhängige Datenpfade
# -------------------------------------------------
BASE_DATA_DIR = "data"

def get_user_data_dir():
    user_ = st.session_state.get("user", "default")
    path = os.path.join(BASE_DATA_DIR, user_)
    os.makedirs(path, exist_ok=True)
    return path

def user_file(name: str) -> str:
    return os.path.join(get_user_data_dir(), name)


# -------------------------------------------------
# ✅ ZENTRALER SPEICHER: dashboard_data.json (pro User)
# -------------------------------------------------
DASHBOARD_JSON = "dashboard_data.json"

DEFAULT_STORE = {
    "klausuren": [],       # list[dict]
    "todos": [],           # list[dict]
    "seminare": [],        # list[dict]
    "lernplan": [],        # list[dict]
    "mood": [],            # list[dict]
    "stundenplan_html": "" # str
}

def _atomic_write_json(path: str, obj: dict):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)

def save_store(store: dict):
    path = user_file(DASHBOARD_JSON)
    _atomic_write_json(path, store)

def normalize_store(data: dict) -> dict:
    """Sorgt dafür, dass alle Keys vorhanden sind und Typen passen."""
    if not isinstance(data, dict):
        return DEFAULT_STORE.copy()

    fixed = DEFAULT_STORE.copy()
    for k in fixed.keys():
        if k in data:
            fixed[k] = data[k]

    if not isinstance(fixed.get("klausuren"), list): fixed["klausuren"] = []
    if not isinstance(fixed.get("todos"), list): fixed["todos"] = []
    if not isinstance(fixed.get("seminare"), list): fixed["seminare"] = []
    if not isinstance(fixed.get("lernplan"), list): fixed["lernplan"] = []
    if not isinstance(fixed.get("mood"), list): fixed["mood"] = []
    if not isinstance(fixed.get("stundenplan_html"), str): fixed["stundenplan_html"] = ""
    return fixed

def load_store() -> dict:
    path = user_file(DASHBOARD_JSON)
    if not os.path.exists(path):
        save_store(DEFAULT_STORE)
        return DEFAULT_STORE.copy()

    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = DEFAULT_STORE.copy()

    return normalize_store(data)

# Store einmal pro Session laden
if "store" not in st.session_state:
    st.session_state["store"] = load_store()

store = st.session_state["store"]


# -------------------------------------------------
# Hilfsfunktionen allgemein
# -------------------------------------------------
def safe_rerun():
    try:
        st.rerun()
    except Exception:
        try:
            st.experimental_rerun()
        except Exception:
            pass

def open_path_or_url(path: str):
    """Ordner oder URL öffnen."""
    path = str(path)

    if path.startswith("http://") or path.startswith("https://"):
        webbrowser.open(path)
        return

    if not os.path.exists(path):
        st.error(f"Pfad nicht gefunden: {path}")
        return

    if sys.platform.startswith("win"):
        os.startfile(path)  # type: ignore
    elif sys.platform.startswith("darwin"):
        subprocess.Popen(["open", path])
    else:
        subprocess.Popen(["xdg-open", path])


# -------------------------------------------------
# Helper: Date parsing/formatting
# -------------------------------------------------
def _to_date_safe(x):
    if x is None or x == "" or (isinstance(x, float) and pd.isna(x)):
        return pd.NaT
    try:
        return pd.to_datetime(x, errors="coerce").date()
    except Exception:
        return pd.NaT

def _date_to_str(d):
    try:
        if pd.isna(d):
            return ""
    except Exception:
        pass
    if isinstance(d, datetime):
        return d.date().isoformat()
    return str(d)


# -------------------------------------------------
# Stundenplan HTML (STORE)
# -------------------------------------------------
def load_stundenplan_html() -> str:
    return store.get("stundenplan_html", "") or ""

def save_stundenplan_html(html: str):
    store["stundenplan_html"] = html
    save_store(store)


# -------------------------------------------------
# Klausuren (STORE)
# -------------------------------------------------
KLAUSUREN_COLS = [
    "fach", "datum", "lernordner", "tage_vorher",
    "archiviert", "note", "ziel_stunden", "gelernt_stunden"
]

def load_klausuren():
    rows = store.get("klausuren", [])
    df = pd.DataFrame(rows)

    for col in KLAUSUREN_COLS:
        if col not in df.columns:
            if col == "tage_vorher":
                df[col] = 21
            elif col == "archiviert":
                df[col] = False
            elif col in ["ziel_stunden", "gelernt_stunden"]:
                df[col] = 0.0
            elif col == "note":
                df[col] = ""
            else:
                df[col] = ""

    df["datum"] = df["datum"].apply(_to_date_safe)
    df["tage_vorher"] = pd.to_numeric(df["tage_vorher"], errors="coerce").fillna(21).astype(int)
    df["archiviert"] = df["archiviert"].astype(bool)
    df["note"] = df["note"].astype(str)
    df["ziel_stunden"] = pd.to_numeric(df["ziel_stunden"], errors="coerce").fillna(0.0)
    df["gelernt_stunden"] = pd.to_numeric(df["gelernt_stunden"], errors="coerce").fillna(0.0)

    return df[KLAUSUREN_COLS].copy()

def save_klausuren(df):
    out = df.copy()
    out["datum"] = out["datum"].apply(_date_to_str)
    out = out[KLAUSUREN_COLS]
    store["klausuren"] = out.to_dict(orient="records")
    save_store(store)

def compute_exam_risk(row, today):
    datum = row["datum"]
    if pd.isna(datum):
        return "unbekannt", "Datum fehlt"

    days_until = (datum - today).days
    if days_until < 0:
        return "vorbei", "Klausur liegt in der Vergangenheit."
    if days_until == 0:
        return "heute", "Heute ist Klausurtag – GO! 🚀"

    ziel = float(row.get("ziel_stunden", 0.0) or 0.0)
    gelernt = float(row.get("gelernt_stunden", 0.0) or 0.0)
    tage_vorher = int(row.get("tage_vorher", 21) or 21)

    if ziel <= 0:
        return "unbekannt", "Keine geplanten Lernstunden hinterlegt."

    progress = gelernt / ziel
    total_window = max(tage_vorher, 1)
    days_elapsed = max(total_window - days_until, 0)
    expected_progress = min(max(days_elapsed / total_window, 0.0), 1.0)

    if progress >= expected_progress * 0.9:
        return "grün", "Du liegst gut im Plan. Weiter so! ✅"
    elif progress >= expected_progress * 0.6:
        return "gelb", "Okay, aber da geht noch was. ⚠️"
    else:
        return "rot", "Rückstand zum Plan – besser Gas geben. ❗"


# -------------------------------------------------
# Todos (STORE)
# -------------------------------------------------
def load_todos():
    data = store.get("todos", [])
    norm = []
    for t in data:
        norm.append(
            {
                "text": t.get("text", ""),
                "done": bool(t.get("done", False)),
                "fach": t.get("fach", ""),
                "wichtig": bool(t.get("wichtig", False)),
                "faellig": t.get("faellig", ""),
            }
        )
    return norm

def save_todos(todos):
    store["todos"] = todos
    save_store(store)


# -------------------------------------------------
# Mood (STORE)
# -------------------------------------------------
MOOD_COLS = ["datum", "stimmung", "stress", "schlaf", "notiz"]

def load_mood():
    rows = store.get("mood", [])
    df = pd.DataFrame(rows)
    for c in MOOD_COLS:
        if c not in df.columns:
            df[c] = "" if c == "notiz" else 0

    df["datum"] = df["datum"].apply(_to_date_safe)
    df["stimmung"] = pd.to_numeric(df["stimmung"], errors="coerce").fillna(0).astype(int)
    df["stress"] = pd.to_numeric(df["stress"], errors="coerce").fillna(0).astype(int)
    df["schlaf"] = pd.to_numeric(df["schlaf"], errors="coerce").fillna(0.0)
    df["notiz"] = df["notiz"].astype(str)
    return df[MOOD_COLS].copy()

def save_mood(df):
    out = df.copy()
    out["datum"] = out["datum"].apply(_date_to_str)
    store["mood"] = out[MOOD_COLS].to_dict(orient="records")
    save_store(store)


# -------------------------------------------------
# Seminare (STORE)
# -------------------------------------------------
SEMINAR_COLS = ["titel", "datum", "uhrzeit1", "datum2", "uhrzeit2", "notiz", "punkte", "absolviert"]

def load_seminare():
    rows = store.get("seminare", [])
    df = pd.DataFrame(rows)

    for col in SEMINAR_COLS:
        if col not in df.columns:
            if col == "punkte":
                df[col] = 0.0
            elif col == "absolviert":
                df[col] = False
            else:
                df[col] = ""

    df["datum"] = df["datum"].apply(_to_date_safe)
    df["datum2"] = df["datum2"].apply(_to_date_safe)
    df["punkte"] = pd.to_numeric(df["punkte"], errors="coerce").fillna(0.0)
    df["absolviert"] = df["absolviert"].astype(bool)
    df["titel"] = df["titel"].astype(str)
    df["uhrzeit1"] = df["uhrzeit1"].astype(str)
    df["uhrzeit2"] = df["uhrzeit2"].astype(str)
    df["notiz"] = df["notiz"].astype(str)
    return df[SEMINAR_COLS].copy()

def save_seminare(df):
    out = df.copy()
    out["datum"] = out["datum"].apply(_date_to_str)
    out["datum2"] = out["datum2"].apply(_date_to_str)
    store["seminare"] = out[SEMINAR_COLS].to_dict(orient="records")
    save_store(store)


# -------------------------------------------------
# Lernplan (STORE)
# -------------------------------------------------
LERNPLAN_COLS = ["fach", "stunden_pro_woche", "priorität"]

def load_lernplan():
    rows = store.get("lernplan", [])
    df = pd.DataFrame(rows)
    for col in LERNPLAN_COLS:
        if col not in df.columns:
            df[col] = "" if col == "fach" else (0.0 if col == "stunden_pro_woche" else 2)

    df["fach"] = df["fach"].astype(str)
    df["stunden_pro_woche"] = pd.to_numeric(df["stunden_pro_woche"], errors="coerce").fillna(0.0)
    df["priorität"] = pd.to_numeric(df["priorität"], errors="coerce").fillna(2).astype(int)
    return df[LERNPLAN_COLS].copy()

def save_lernplan(df):
    out = df.copy()
    store["lernplan"] = out[LERNPLAN_COLS].to_dict(orient="records")
    save_store(store)


# -------------------------------------------------
# Datei-Extraktion für Lernzettel
# -------------------------------------------------
def extract_text_from_file(uploaded_file):
    filename = uploaded_file.name.lower()

    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if filename.endswith(".docx"):
        try:
            file_bytes = uploaded_file.read()
            doc = Document(BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            return f"(Fehler beim Lesen der Word-Datei: {e})"

    if filename.endswith(".pdf"):
        try:
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception:
            return "(PDF konnte nicht gelesen werden)"

    return "(Dateiformat nicht unterstützt)"


# -------------------------------------------------
# Streamlit Setup
# -------------------------------------------------
st.set_page_config(page_title="Uni-Dashboard", page_icon="📚", layout="wide")

st.sidebar.title("📚 Uni-Dashboard (v5)")
page = st.sidebar.radio(
    "Bereich wählen",
    [
        "Tagesübersicht",
        "Stundenplan",
        "Klausuren & Lernen",
        "To-Do & Hausaufgaben",
        "Seminare & Punkte",
        "Lernplan Woche",
        "Lernzettel erstellen",
        "PDFs zusammenfügen",
        "PDF erstellen",
        "LaTeX",
        "Mood-Tracker & Stressradar",
    ],
)

# ✅ UPGRADE: Backup/Restore in Sidebar
st.sidebar.divider()
st.sidebar.subheader("💾 Backup / Restore")

backup_json = json.dumps(store, ensure_ascii=False, indent=2).encode("utf-8")
st.sidebar.download_button(
    "⬇️ Backup herunterladen (JSON)",
    data=backup_json,
    file_name=f"dashboard_backup_{st.session_state.get('user','user')}.json",
    mime="application/json",
    use_container_width=True,
)

uploaded_backup = st.sidebar.file_uploader(
    "⬆️ Backup wiederherstellen (JSON)",
    type=["json"],
    help="Lädt ein Backup und überschreibt deine aktuellen Daten.",
)

if uploaded_backup is not None:
    try:
        imported = json.loads(uploaded_backup.read().decode("utf-8"))
        imported = normalize_store(imported)

        st.sidebar.warning("Achtung: Restore überschreibt ALLE aktuellen Daten.")
        if st.sidebar.button("✅ Restore jetzt durchführen", use_container_width=True):
            st.session_state["store"] = imported
            store.clear()
            store.update(imported)
            save_store(store)
            st.sidebar.success("Restore erfolgreich ✅")
            safe_rerun()

    except Exception as e:
        st.sidebar.error(f"Backup konnte nicht geladen werden: {e}")

today = datetime.today().date()

# ✅ Alles aus dashboard_data.json laden
klausuren = load_klausuren()
todos = load_todos()
seminare = load_seminare()
lernplan = load_lernplan()


# -------------------------------------------------
# 0️⃣ TAGESÜBERSICHT
# -------------------------------------------------
if page == "Tagesübersicht":
    st.title("🏠 Tagesübersicht")
    st.subheader(f"Heute: {today.strftime('%A, %d.%m.%Y')}")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### 📆 Nächste Klausuren")
        aktive = klausuren[~klausuren["archiviert"]].copy()
        aktive = aktive[pd.notna(aktive["datum"])]
        aktive = aktive[aktive["datum"] >= today].sort_values("datum")

        if aktive.empty:
            st.write("Keine anstehenden Klausuren 🙌")
        else:
            for _, row in aktive.head(3).iterrows():
                days_until = (row["datum"] - today).days
                risk, msg = compute_exam_risk(row, today)
                st.write(f"**{row['fach']}** – in {days_until} Tag(en) ({row['datum'].strftime('%d.%m.%Y')})")
                if risk == "grün":
                    st.success(msg)
                elif risk == "gelb":
                    st.warning(msg)
                elif risk == "rot":
                    st.error(msg)
                else:
                    st.info(msg)

    with col2:
        st.markdown("### ✅ Wichtige To-Dos (Top 5)")
        open_todos = [t for t in todos if not t["done"]]
        if not open_todos:
            st.write("Alles erledigt, stark! 🎉")
        else:
            def todo_sort_key(t):
                due = t.get("faellig") or ""
                try:
                    d = datetime.fromisoformat(due).date()
                except Exception:
                    d = today + timedelta(days=365)
                return (not t.get("wichtig", False), d)

            open_todos_sorted = sorted(open_todos, key=todo_sort_key)[:5]
            for t in open_todos_sorted:
                label = t["text"]
                if t.get("fach"):
                    label += f" ({t['fach']})"
                if t.get("faellig"):
                    try:
                        d = datetime.fromisoformat(t["faellig"]).date()
                        label += f" – bis {d.strftime('%d.%m.%Y')}"
                    except Exception:
                        label += f" – bis {t['faellig']}"
                st.write(("🔴 " if t.get("wichtig") else "🟢 ") + label)

    with col3:
        st.markdown("### 🎓 Seminare & Stimmung")
        sem_today = seminare[pd.notna(seminare["datum"])]
        sem_today = sem_today[sem_today["datum"] == today]
        if not sem_today.empty:
            st.write("**Heutige Seminare:**")
            for _, row in sem_today.iterrows():
                info = row["titel"]
                if row.get("uhrzeit1", ""):
                    info += f" – {row['uhrzeit1']}"
                if row.get("notiz", "").strip():
                    info += f" ({row['notiz']})"
                info += f" – {row['punkte']} Punkte"
                st.write(f"- {info}")
        else:
            sem_next = seminare[pd.notna(seminare["datum"])]
            sem_next = sem_next[sem_next["datum"] > today].sort_values("datum")
            if not sem_next.empty:
                nxt = sem_next.iloc[0]
                text = f"{nxt['titel']} am {nxt['datum'].strftime('%d.%m.%Y')}"
                if nxt.get("uhrzeit1", ""):
                    text += f", {nxt['uhrzeit1']}"
                if nxt.get("notiz", "").strip():
                    text += f" ({nxt['notiz']})"
                text += f" – {nxt['punkte']} Punkte"
                st.write("**Nächstes Seminar:**")
                st.write(text)
            else:
                st.write("Keine Seminare eingetragen.")

        mood_df = load_mood()
        if not mood_df.empty:
            last = mood_df.sort_values("datum").iloc[-1]
            st.write("---")
            st.write("**Letzter Stimmungseintrag:**")
            st.write(f"Stimmung: {last['stimmung']}/10")
            st.write(f"Stress: {last['stress']}/10")
            st.write(f"Schlaf: {last['schlaf']} h")
        else:
            st.write("Noch kein Mood-Tracking gestartet.")

    st.markdown("---")
    st.subheader("⏱️ Lernzeit-Timer (Pomodoro light)")

    if "timer_mode" not in st.session_state:
        st.session_state["timer_mode"] = None
    if "timer_start" not in st.session_state:
        st.session_state["timer_start"] = None
    if "timer_duration" not in st.session_state:
        st.session_state["timer_duration"] = 0
    if "timer_learn_minutes" not in st.session_state:
        st.session_state["timer_learn_minutes"] = 25
    if "timer_break_minutes" not in st.session_state:
        st.session_state["timer_break_minutes"] = 5
    if "timer_sound_played" not in st.session_state:
        st.session_state["timer_sound_played"] = False
    if "timer_exam_index" not in st.session_state:
        st.session_state["timer_exam_index"] = None
    if "timer_logged_to_exam" not in st.session_state:
        st.session_state["timer_logged_to_exam"] = False

    sound_dir = "sounds"
    available_sounds = []
    if os.path.isdir(sound_dir):
        for f in os.listdir(sound_dir):
            if f.lower().endswith((".mp3", ".wav", ".ogg")):
                available_sounds.append(f)

    if "timer_sound_file" not in st.session_state:
        st.session_state["timer_sound_file"] = available_sounds[0] if available_sounds else None

    col_t1, col_t2, col_t3 = st.columns(3)
    with col_t1:
        st.number_input("Lernphase (Minuten)", min_value=5, max_value=180, key="timer_learn_minutes")
    with col_t2:
        st.number_input("Pause (Minuten)", min_value=1, max_value=60, key="timer_break_minutes")

    st.markdown("### 🔊 Sound-Einstellungen")
    col_s1, col_s2 = st.columns([2, 1])
    with col_s1:
        options = ["(kein Sound)"] + available_sounds if available_sounds else ["(kein Sound)"]
        current = st.session_state.get("timer_sound_file")
        default_index = options.index(current) if current in options else 0
        choice = st.selectbox("Alarm-Sound", options, index=default_index)
        st.session_state["timer_sound_file"] = None if choice == "(kein Sound)" else choice

    with col_s2:
        if st.button("Sound testen"):
            sound_name = st.session_state.get("timer_sound_file")
            if sound_name:
                try:
                    with open(os.path.join(sound_dir, sound_name), "rb") as f:
                        audio_bytes = f.read()
                    st.audio(audio_bytes, format="audio/mp3")
                except FileNotFoundError:
                    st.warning(f"Sounddatei '{sound_name}' wurde nicht gefunden.")
            else:
                st.info("Kein Sound ausgewählt.")

    aktive_klausuren = klausuren[~klausuren["archiviert"]]
    exam_options = {"(keine Verknüpfung)": None}
    for idx, row in aktive_klausuren.iterrows():
        label = f"{row['fach']} – {row['datum'].strftime('%d.%m.%Y')}" if pd.notna(row["datum"]) else f"{row['fach']} – (ohne Datum)"
        exam_options[label] = idx

    selected_label = st.selectbox("Timer mit Klausur verknüpfen (optional)", list(exam_options.keys()))
    st.session_state["timer_exam_index"] = exam_options[selected_label]

    def start_timer(mode):
        minutes = st.session_state["timer_learn_minutes"] if mode == "Lernphase" else st.session_state["timer_break_minutes"]
        st.session_state["timer_mode"] = mode
        st.session_state["timer_start"] = datetime.now().isoformat()
        st.session_state["timer_duration"] = int(minutes * 60)
        st.session_state["timer_sound_played"] = False
        st.session_state["timer_logged_to_exam"] = False

    with col_t3:
        if st.button("Lernphase starten"):
            start_timer("Lernphase")
        if st.button("Pause starten"):
            start_timer("Pause")
        if st.button("Timer zurücksetzen"):
            st.session_state["timer_mode"] = None
            st.session_state["timer_start"] = None
            st.session_state["timer_duration"] = 0
            st.session_state["timer_sound_played"] = False
            st.session_state["timer_logged_to_exam"] = False

    st.write("---")
    if st.session_state["timer_mode"] and st.session_state["timer_start"]:
        mode = st.session_state["timer_mode"]
        start_dt = datetime.fromisoformat(st.session_state["timer_start"])
        duration = st.session_state["timer_duration"]
        elapsed = (datetime.now() - start_dt).total_seconds()
        remaining = max(duration - elapsed, 0)
        progress = min(max(elapsed / duration, 0), 1) if duration > 0 else 0

        mins = int(remaining // 60)
        secs = int(remaining % 60)

        st.write(f"Aktiver Timer: **{mode}**")
        st.progress(progress)

        if remaining > 0:
            st.write(f"Noch {mins:02d}:{secs:02d} Minuten")
        else:
            st.success("Fertig! ✅" if mode == "Lernphase" else "Pause vorbei! 💪")

            if mode == "Lernphase" and not st.session_state["timer_logged_to_exam"]:
                exam_idx = st.session_state.get("timer_exam_index")
                if exam_idx is not None and exam_idx in klausuren.index:
                    minutes = st.session_state["timer_learn_minutes"]
                    hours = minutes / 60.0
                    vorher = float(klausuren.at[exam_idx, "gelernt_stunden"])
                    klausuren.at[exam_idx, "gelernt_stunden"] = vorher + hours
                    save_klausuren(klausuren)
                    st.success(f"{hours:.2f} h wurden für '{klausuren.at[exam_idx, 'fach']}' gutgeschrieben.")
                st.session_state["timer_logged_to_exam"] = True

            if not st.session_state["timer_sound_played"]:
                sound_name = st.session_state.get("timer_sound_file")
                if sound_name:
                    try:
                        with open(os.path.join(sound_dir, sound_name), "rb") as f:
                            audio = f.read()
                        st.audio(audio, format="audio/mp3")
                    except FileNotFoundError:
                        st.warning(f"Sounddatei '{sound_name}' wurde nicht gefunden.")
                st.session_state["timer_sound_played"] = True
    else:
        st.write("Kein Timer aktiv. Starte eine Lernphase oder Pause.")


# -------------------------------------------------
# 1️⃣ STUNDENPLAN – HTML (JETZT IN JSON)
# -------------------------------------------------
elif page == "Stundenplan":
    st.title("📅 Stundenplan (HTML)")

    st.markdown(
        "✅ Speichert jetzt in **dashboard_data.json** (pro User).\n\n"
        "- Scrollbalken nach unten/rechts\n"
        "- Upload → Vorschau → Speichern"
    )

    FRAME_HEIGHT = 800
    FRAME_WIDTH = 1200

    html_content = load_stundenplan_html().strip()

    if html_content:
        st.markdown("### 🔍 Aktuell gespeicherter Stundenplan")
        components.html(html_content, height=FRAME_HEIGHT, width=FRAME_WIDTH, scrolling=True)
    else:
        st.info("Noch kein Stundenplan gespeichert.")

    st.markdown("---")
    st.subheader("📤 Neuen HTML-Stundenplan hochladen")

    uploaded_html = st.file_uploader("HTML-Datei auswählen:", type=["html", "htm"])

    if uploaded_html is not None:
        if (
            "stundenplan_html_upload" not in st.session_state
            or st.session_state.get("stundenplan_html_upload_name") != uploaded_html.name
        ):
            html_text = uploaded_html.read().decode("utf-8", errors="ignore")
            st.session_state["stundenplan_html_upload"] = html_text
            st.session_state["stundenplan_html_upload_name"] = uploaded_html.name

        html_upload_content = st.session_state["stundenplan_html_upload"]

        st.success(f"Neue HTML-Datei `{uploaded_html.name}` geladen ✅")
        st.markdown("### 🧾 Vorschau")
        components.html(html_upload_content, height=FRAME_HEIGHT, width=FRAME_WIDTH, scrolling=True)

        if st.button("💾 Diesen Stundenplan für meinen Account speichern"):
            save_stundenplan_html(html_upload_content)
            st.success("Stundenplan gespeichert ✅")
            safe_rerun()
    else:
        st.info("Lade eine HTML-Datei hoch, um sie anzuschauen oder zu speichern.")


# -------------------------------------------------
# 2️⃣ KLAUSUREN & LERNEN
# -------------------------------------------------
elif page == "Klausuren & Lernen":
    st.title("📝 Klausuren & Lernen")

    view = st.radio("Ansicht", ["Aktive Klausuren", "Archiv"])
    df_view = klausuren[~klausuren["archiviert"]] if view == "Aktive Klausuren" else klausuren[klausuren["archiviert"]]

    if df_view.empty:
        st.info("Keine Klausuren in dieser Ansicht.")
    else:
        for idx, row in df_view.sort_values("datum", na_position="last").iterrows():
            st.markdown("---")
            col1, col2 = st.columns([2, 1])

            with col1:
                st.subheader(f"📌 {row['fach']}")
                if pd.notna(row["datum"]):
                    st.write(f"**Datum:** {row['datum'].strftime('%d.%m.%Y')}")
                    days_until = (row["datum"] - today).days
                    st.write(f"**Noch:** {days_until} Tag(e)")
                else:
                    st.write("**Datum:** -")

                if not row["archiviert"]:
                    ziel = st.number_input(
                        "Geplante Lernstunden insgesamt",
                        min_value=0.0, max_value=500.0, step=0.5,
                        value=float(row.get("ziel_stunden", 0.0) or 0.0),
                        key=f"ziel_{idx}",
                    )
                    gelernt = st.number_input(
                        "Bisher gelernte Stunden",
                        min_value=0.0, max_value=500.0, step=0.5,
                        value=float(row.get("gelernt_stunden", 0.0) or 0.0),
                        key=f"gelernt_{idx}",
                    )
                    klausuren.at[idx, "ziel_stunden"] = ziel
                    klausuren.at[idx, "gelernt_stunden"] = gelernt

                    if ziel > 0:
                        progress = max(min(gelernt / ziel, 1.0), 0.0)
                        st.write("**Lernfortschritt:**")
                        st.progress(progress)
                        st.write(f"{gelernt:.1f} / {ziel:.1f} Stunden")

                        risk, msg = compute_exam_risk(klausuren.loc[idx], today)
                        if risk == "grün":
                            st.success(msg)
                        elif risk == "gelb":
                            st.warning(msg)
                        elif risk == "rot":
                            st.error(msg)
                        else:
                            st.info(msg)
                    else:
                        st.info("Noch keine geplanten Lernstunden hinterlegt.")

                    new_tage = st.number_input(
                        "Empfohlene Tage vorher zu lernen",
                        min_value=1, max_value=180,
                        value=int(row["tage_vorher"]),
                        key=f"tage_{idx}",
                    )
                    klausuren.at[idx, "tage_vorher"] = new_tage
                else:
                    try:
                        default_note = float(row.get("note", "0") or 0)
                    except ValueError:
                        default_note = 0.0
                    note = st.number_input(
                        "Note (0–15):",
                        min_value=0.0, max_value=15.0,
                        value=default_note, step=0.5,
                        key=f"note_{idx}",
                    )
                    klausuren.at[idx, "note"] = str(note)

                    if note > 4:
                        st.success("Bestanden 🎉")
                    else:
                        st.error("Nicht bestanden ❌")

            with col2:
                st.write("**Aktionen:**")
                if st.button("Ordner öffnen", key=f"ordner_{idx}"):
                    open_path_or_url(row["lernordner"])

                if not row["archiviert"]:
                    if st.button("Archivieren", key=f"archiv_{idx}"):
                        klausuren.at[idx, "archiviert"] = True
                        save_klausuren(klausuren)
                        safe_rerun()
                else:
                    if st.button("Löschen", key=f"del_{idx}"):
                        klausuren = klausuren.drop(idx).reset_index(drop=True)
                        save_klausuren(klausuren)
                        safe_rerun()

        save_klausuren(klausuren)

    st.markdown("---")
    st.subheader("➕ Neue Klausur")

    new_fach = st.text_input("Fach")
    new_datum = st.date_input("Datum", value=today)
    new_ordner = st.text_input("Lernordner")
    new_tage = st.number_input("Tage vorher", min_value=1, max_value=180, value=21)
    new_ziel = st.number_input("Geplante Lernstunden (optional)", min_value=0.0, max_value=500.0, step=0.5, value=0.0)

    if st.button("Klausur speichern"):
        klausuren.loc[len(klausuren)] = [
            new_fach, new_datum, new_ordner, int(new_tage),
            False, "", float(new_ziel), 0.0
        ]
        save_klausuren(klausuren)
        st.success("Klausur wurde hinzugefügt!")
        safe_rerun()


# -------------------------------------------------
# 3️⃣ TO-DO LISTE
# -------------------------------------------------
elif page == "To-Do & Hausaufgaben":
    st.title("📋 To-Do Liste")

    todos = load_todos()
    delete_index = None

    for idx, todo in enumerate(todos):
        col1, col2, col3 = st.columns([0.6, 0.2, 0.2])

        with col1:
            done = st.checkbox(todo["text"], todo["done"], key=f"done_{idx}")
            todos[idx]["done"] = done

        with col2:
            wichtig = st.checkbox("Wichtig", todo["wichtig"], key=f"wicht_{idx}")
            todos[idx]["wichtig"] = wichtig

        with col3:
            if st.button("🗑️", key=f"delete_{idx}"):
                delete_index = idx

    if delete_index is not None:
        todos.pop(delete_index)

    save_todos(todos)

    st.markdown("---")
    st.subheader("➕ Neue Aufgabe")

    new_text = st.text_input("Aufgabe")
    new_fach = st.text_input("Fach")
    new_due = st.date_input("Fällig bis", value=today)

    if st.button("Aufgabe hinzufügen"):
        todos.append(
            {
                "text": new_text,
                "done": False,
                "fach": new_fach,
                "wichtig": False,
                "faellig": str(new_due),
            }
        )
        save_todos(todos)
        st.success("Aufgabe hinzugefügt!")
        safe_rerun()


# -------------------------------------------------
# 4️⃣ SEMINARE & PUNKTE
# -------------------------------------------------
elif page == "Seminare & Punkte":
    st.title("🎓 Seminare & Punkte")

    if seminare.empty:
        total_all = 0.0
        total_done = 0.0
    else:
        total_all = seminare["punkte"].sum()
        total_done = seminare.loc[seminare["absolviert"], "punkte"].sum()

    col_a, col_b = st.columns(2)
    with col_a:
        st.metric("Gesamtpunkte (alle Seminare)", f"{total_all:.1f}")
    with col_b:
        st.metric("Gesammelte Punkte (absolvierte Seminare)", f"{total_done:.1f}")

    st.markdown("---")
    st.subheader("📋 Seminarübersicht")

    if seminare.empty:
        st.info("Trage unten dein erstes Seminar ein.")
    else:
        delete_idx = None

        for idx, row in seminare.sort_values("datum", na_position="last").iterrows():
            st.markdown("---")
            c1, c2, c3, c4 = st.columns([2, 1, 1, 0.5])

            with c1:
                st.write(f"**Titel:** {row['titel']}")

                if pd.notna(row["datum"]):
                    line1 = row["datum"].strftime("%d.%m.%Y")
                    if row.get("uhrzeit1", ""):
                        line1 += f", {row['uhrzeit1']}"
                    st.write(f"**Termin 1:** {line1}")
                else:
                    st.write("**Termin 1:** -")

                datum2_val = row.get("datum2", pd.NaT)
                if pd.notna(datum2_val):
                    d2_str = datum2_val.strftime("%d.%m.%Y")
                    if row.get("uhrzeit2", ""):
                        d2_str += f", {row['uhrzeit2']}"
                    st.write(f"**Termin 2:** {d2_str}")

                notiz_str = str(row.get("notiz", "") or "").strip()
                st.write(f"**Notiz:** {notiz_str if notiz_str else '-'}")

            with c2:
                punkte_val = st.number_input(
                    "Punkte",
                    min_value=0.0, max_value=30.0, step=0.5,
                    value=float(row["punkte"]),
                    key=f"sem_punkte_{idx}",
                )
                seminare.at[idx, "punkte"] = punkte_val

            with c3:
                absolviert_val = st.checkbox(
                    "Absolviert?",
                    value=bool(row["absolviert"]),
                    key=f"sem_done_{idx}",
                )
                seminare.at[idx, "absolviert"] = absolviert_val

            with c4:
                if st.button("🗑️", key=f"sem_del_{idx}"):
                    delete_idx = idx

        if delete_idx is not None:
            seminare = seminare.drop(delete_idx).reset_index(drop=True)

        save_seminare(seminare)

    st.markdown("---")
    st.subheader("➕ Neues Seminar hinzufügen")

    new_titel = st.text_input("Titel des Seminars")

    col_d1, col_d2 = st.columns(2)
    with col_d1:
        new_datum = st.date_input("Datum – Termin 1", value=today)
    with col_d2:
        new_uhrzeit1 = st.text_input("Uhrzeit – Termin 1 (z.B. 10:00–12:00)", value="")

    second_day = st.checkbox("Seminar hat einen zweiten Termin?", value=False)

    new_datum2 = pd.NaT
    new_uhrzeit2 = ""
    if second_day:
        col_d3, col_d4 = st.columns(2)
        with col_d3:
            new_datum2 = st.date_input("Datum – Termin 2", value=today)
        with col_d4:
            new_uhrzeit2 = st.text_input("Uhrzeit – Termin 2", value="")

    new_notiz = st.text_area("Notiz (Raum/Link/Anbieter/…)", value="")

    new_punkte = st.number_input("Punkte", min_value=0.0, max_value=30.0, step=0.5, value=0.0)
    new_done = st.checkbox("Bereits absolviert?", value=False)

    if st.button("Seminar speichern"):
        if not new_titel.strip():
            st.warning("Bitte einen Seminartitel eingeben.")
        else:
            new_row = {
                "titel": new_titel.strip(),
                "datum": new_datum,
                "uhrzeit1": new_uhrzeit1.strip(),
                "datum2": new_datum2 if second_day else pd.NaT,
                "uhrzeit2": new_uhrzeit2.strip() if second_day else "",
                "notiz": new_notiz.strip(),
                "punkte": float(new_punkte),
                "absolviert": bool(new_done),
            }
            seminare = pd.concat([seminare, pd.DataFrame([new_row])], ignore_index=True)
            save_seminare(seminare)
            st.success("Seminar hinzugefügt.")
            safe_rerun()


# -------------------------------------------------
# 5️⃣ LERNPLAN WOCHE
# -------------------------------------------------
elif page == "Lernplan Woche":
    st.title("📆 Lernplan für die Woche")

    if lernplan.empty:
        st.info("Noch kein Lernplan angelegt. Füge unten Fächer hinzu.")
    else:
        st.subheader("📚 Übersicht Lernfächer")

        delete_idx = None
        for idx, row in lernplan.iterrows():
            c1, c2, c3, c4 = st.columns([2, 1, 1, 0.5])
            with c1:
                fach_val = st.text_input("Fach", value=row["fach"], key=f"lp_fach_{idx}")
                lernplan.at[idx, "fach"] = fach_val
            with c2:
                stunden_val = st.number_input(
                    "Stunden/Woche",
                    min_value=0.0, max_value=50.0, step=0.5,
                    value=float(row["stunden_pro_woche"]),
                    key=f"lp_stunden_{idx}",
                )
                lernplan.at[idx, "stunden_pro_woche"] = stunden_val
            with c3:
                prio_val = st.selectbox(
                    "Priorität (1=hoch,3=niedrig)",
                    [1, 2, 3],
                    index={1: 0, 2: 1, 3: 2}[int(row["priorität"]) if row["priorität"] in [1, 2, 3] else 1],
                    key=f"lp_prio_{idx}",
                )
                lernplan.at[idx, "priorität"] = prio_val
            with c4:
                if st.button("🗑️", key=f"lp_del_{idx}"):
                    delete_idx = idx

        if delete_idx is not None:
            lernplan = lernplan.drop(delete_idx).reset_index(drop=True)

        save_lernplan(lernplan)

        st.markdown("---")
        st.subheader("📊 Geplante Gesamtstunden pro Woche")
        total_hours = lernplan["stunden_pro_woche"].sum()
        st.write(f"**Summe:** {total_hours:.1f} Stunden/Woche")

        st.markdown("### Vorschlag: Verteilung auf die Wochentage (Mo–Fr)")
        days = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag", "Samstag", "Sonntag"]
        plan_dict = {d: 0.0 for d in days}

        for _, row in lernplan.iterrows():
            h = float(row["stunden_pro_woche"])
            if h <= 0:
                continue
            share = h / 5.0
            for d in days[:5]:
                plan_dict[d] += share

        plan_df = pd.DataFrame({"Tag": list(plan_dict.keys()), "Geplante Lernstunden": list(plan_dict.values())})
        plan_df["Geplante Lernstunden"] = plan_df["Geplante Lernstunden"].map(lambda x: f"{x:.1f} h")
        st.table(plan_df)

    st.markdown("---")
    st.subheader("➕ Neues Fach zum Lernplan hinzufügen")

    lp_fach = st.text_input("Fachname")
    lp_stunden = st.number_input("Stunden pro Woche", min_value=0.0, max_value=50.0, step=0.5, value=0.0)
    lp_prio = st.selectbox("Priorität", [1, 2, 3], help="1 = sehr wichtig, 3 = weniger wichtig")

    if st.button("Fach zum Lernplan hinzufügen"):
        if not lp_fach.strip():
            st.warning("Bitte Fachname eintragen.")
        else:
            new_row = {"fach": lp_fach.strip(), "stunden_pro_woche": float(lp_stunden), "priorität": int(lp_prio)}
            lernplan = pd.concat([lernplan, pd.DataFrame([new_row])], ignore_index=True)
            save_lernplan(lernplan)
            st.success("Fach zum Lernplan hinzugefügt.")
            safe_rerun()


# -------------------------------------------------
# 6️⃣ LERNZETTEL ERSTELLEN
# -------------------------------------------------
elif page == "Lernzettel erstellen":
    st.title("🧠 Lernzettel erstellen")

    uploaded_files = st.file_uploader(
        "Dateien hochladen (PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        st.info(f"{len(uploaded_files)} Datei(en) ausgewählt.")
        if st.button("📘 Dokumente zusammenführen"):
            combined = ""
            for uf in uploaded_files:
                text = extract_text_from_file(uf)
                combined += f"\n\n##### Datei: {uf.name} #####\n\n{text}"
            st.session_state["combined_text"] = combined

    if "combined_text" in st.session_state:
        st.subheader("📄 Zusammengeführtes Dokument")
        st.info("Bearbeite den Text, bevor du ihn an eine KI weitergibst.")

        edited = st.text_area("Dokument bearbeiten:", st.session_state["combined_text"], height=350)
        st.session_state["combined_text"] = edited

        st.markdown("---")
        st.subheader("🤖 Mit KI weiterarbeiten")

        col1, col2, col3 = st.columns(3)
        with col1:
            st.link_button("➡️ ChatGPT", "https://chat.openai.com/")
        with col2:
            st.link_button("➡️ DeepSeek", "https://chat.deepseek.com/")
        with col3:
            st.link_button("➡️ Gemini", "https://gemini.google.com/app")

        st.info(
            'Tipp: Schreibe z. B.: "Bitte strukturiere den Text als Lernzettel mit '
            'Definitionen, Beispielen und Eselsbrücken."'
        )

        if st.button("📥 Dokument als Word (.docx) speichern"):
            doc = Document()
            for line in st.session_state["combined_text"].split("\n"):
                doc.add_paragraph(line)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                "📄 Word herunterladen",
                buffer,
                file_name="lernzettel.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# -------------------------------------------------
# 7️⃣ PDFs ZUSAMMENFÜGEN
# -------------------------------------------------
elif page == "PDFs zusammenfügen":
    st.title("📚 PDFs zusammenfügen")

    uploaded_pdfs = st.file_uploader(
        "Mehrere PDF-Dateien auswählen, die zu einer zusammengefügt werden sollen:",
        type=["pdf"],
        accept_multiple_files=True,
    )

    if uploaded_pdfs:
        st.info(f"{len(uploaded_pdfs)} PDF-Datei(en) ausgewählt.")

        if st.button("📎 PDFs zu einer Datei zusammenfügen"):
            merger = PyPDF2.PdfMerger()
            for pdf_file in uploaded_pdfs:
                try:
                    merger.append(pdf_file)
                except Exception:
                    st.error(f"Fehler beim Verarbeiten von {pdf_file.name}")
            out_buffer = BytesIO()
            merger.write(out_buffer)
            merger.close()
            out_buffer.seek(0)

            st.success("PDFs wurden erfolgreich zusammengefügt.")
            st.download_button(
                "📄 Zusammengeführte PDF herunterladen",
                out_buffer,
                file_name="zusammengefuegt.pdf",
                mime="application/pdf",
            )
    else:
        st.info("Bitte wähle mindestens zwei PDFs aus.")


# -------------------------------------------------
# 7.5 🧾 PDF ERSTELLEN (externes Tool)
# -------------------------------------------------
elif page == "PDF erstellen":
    st.title("🧾 PDF erstellen")

    st.write(
        "Zum Erstellen, Konvertieren und Bearbeiten von PDFs nutzt dieses Dashboard "
        "das externe Tool **PDF24**."
    )

    st.link_button("➡️ PDF24 öffnen", "https://tools.pdf24.org/de/")

    st.caption(
        "Der Link öffnet sich in einem neuen Tab. Dort kannst du PDFs z.B. erstellen, "
        "zusammenfügen, komprimieren oder in andere Formate umwandeln."
    )


# -------------------------------------------------
# 9️⃣ LaTeX – (dein Original-Teil bleibt)
# -------------------------------------------------
elif page == "LaTeX":
    st.title("📐 LaTeX – kurz erklärt & Formelsammlung")

    st.subheader("👀 Was ist LaTeX?")
    st.markdown(
        r"""
LaTeX ist ein Textsatzsystem für **wissenschaftliche Arbeiten** (z.B. Bachelorarbeit, Hausarbeiten, Paper).

Statt mit der Maus zu formatieren (wie in Word), schreibst du **Befehle im Text**, z.B.:

Inline-Formel im Text: `$E = mc^2$`  
Zentrierte Formel: `\[ E = mc^2 \]`

LaTeX ist besonders stark, wenn du:
- viele **Formeln** hast
- ein sauberes **Inhaltsverzeichnis** brauchst
- automatisch ein **Literaturverzeichnis** erzeugen willst
"""
    )

    st.markdown("---")
    st.subheader("📝 LaTeX für Text, Titel & Listen")

    text_snippets = [
        {
            "title": "Dokument mit Titel",
            "latex": r"""\documentclass[a4paper,12pt]{article}

\title{Titel der Arbeit}
\author{Dein Name}
\date{\today}

\begin{document}
\maketitle

Hier beginnt dein Text.

\end{document}""",
            "desc": "Einfache Grundstruktur eines Dokuments mit Titelblatt.",
        },
        {
            "title": "Abschnitt & Unterabschnitt",
            "latex": r"""\section{Einleitung}
\subsection{Motivation}
Dies ist ein normaler Fließtext in LaTeX.""",
            "desc": "Überschriften für Kapitel und Unterkapitel.",
        },
        {
            "title": "Fett & kursiv",
            "latex": r"""Dies ist \textbf{fetter Text} und dies ist \textit{kursiver Text}.""",
            "desc": "Hervorhebung im Fließtext.",
        },
        {
            "title": "Aufzählung (Liste)",
            "latex": r"""\begin{itemize}
  \item Erster Punkt
  \item Zweiter Punkt
  \item Dritter Punkt
\end{itemize}""",
            "desc": "Unsortierte Liste mit Punkten.",
        },
        {
            "title": "Nummerierte Liste",
            "latex": r"""\begin{enumerate}
  \item Erster Punkt
  \item Zweiter Punkt
  \item Dritter Punkt
\end{enumerate}""",
            "desc": "Sortierte Liste mit Nummerierung.",
        },
        {
            "title": "Zitat / Zitat-Umgebung",
            "latex": r"""\begin{quote}
Dies ist ein eingerücktes Zitat.
\end{quote}""",
            "desc": "Zitat oder wichtige Textpassage hervorheben.",
        },
        {
            "title": "Mathe-Umgebung im Fließtext",
            "latex": r"""Dies ist eine Formel im Text: $E = mc^2$.""",
            "desc": "Inline-Math mit Dollarzeichen.",
        },
        {
            "title": "Zentrierte Formel",
            "latex": r"""\[
E = mc^2
\]""",
            "desc": "Formel in einer eigenen zentrierten Zeile.",
        },
    ]

    for ex in text_snippets:
        st.markdown("---")
        st.markdown(f"### {ex['title']}")
        st.caption(ex["desc"])
        st.code(ex["latex"], language="latex")

    st.markdown("---")
    st.subheader("🧮 Kleine LaTeX-Formelsammlung")

    formulas = [
        {"title": "Bruch", "latex": r"\frac{a}{b}", "desc": "Ein einfacher Bruch a/b"},
        {"title": "Potenzen", "latex": r"a^2,\; a^n", "desc": "Quadrat und allgemeine Potenz"},
        {"title": "Wurzel", "latex": r"\sqrt{a},\; \sqrt[n]{a}", "desc": "Quadratwurzel und n-te Wurzel"},
        {"title": "Summenzeichen", "latex": r"\sum_{i=1}^{n} i", "desc": "Summe von i = 1 bis n"},
        {"title": "Produktzeichen", "latex": r"\prod_{i=1}^{n} a_i", "desc": "Produkt über n Faktoren"},
        {"title": "Mitternachtsformel", "latex": r"\frac{-b \pm \sqrt{b^2 - 4ac}}{2a}", "desc": "Quadratische Gleichung"},
        {"title": "Ableitung", "latex": r"\frac{d}{dx} f(x)", "desc": "Ableitung"},
        {"title": "Integral", "latex": r"\int_{a}^{b} f(x)\,dx", "desc": "Bestimmtes Integral"},
        {"title": "Grenzwert", "latex": r"\lim_{x \to \infty} f(x)", "desc": "Grenzwert"},
        {"title": "Matrix 2×2", "latex": r"\begin{pmatrix} a & b \\ c & d \end{pmatrix}", "desc": "2×2 Matrix"},
    ]

    for f in formulas:
        st.markdown("---")
        st.markdown(f"### {f['title']}")
        st.latex(f["latex"])
        st.caption(f["desc"])
        st.markdown("LaTeX-Code:")
        st.code(f["latex"], language="latex")

    st.markdown("---")
    st.info("Tipp: Inline-Formeln: `$ ... $`, Blockformeln: `\\[ ... \\]`")


# -------------------------------------------------
# 8️⃣ MOOD-TRACKER & STRESSRADAR (STORE)
# -------------------------------------------------
elif page == "Mood-Tracker & Stressradar":
    st.title("🌟 Mood-Tracker & Stressradar")

    mood_df = load_mood()

    st.subheader("📅 Heutiger Eintrag")

    col1, col2, col3 = st.columns(3)
    with col1:
        datum = st.date_input("Datum", value=today)
    with col2:
        stimmung = st.slider("Stimmung (1 = schlecht, 10 = top)", 1, 10, 7)
    with col3:
        stress = st.slider("Stresslevel (1 = entspannt, 10 = extrem)", 1, 10, 5)

    schlaf = st.slider("Schlaf letzte Nacht (Stunden)", 0.0, 12.0, 7.0, 0.5)
    notiz = st.text_area("Notiz (optional)", "")

    if st.button("Eintrag speichern"):
        new_row = {"datum": datum, "stimmung": stimmung, "stress": stress, "schlaf": schlaf, "notiz": notiz}
        mood_df = pd.concat([mood_df, pd.DataFrame([new_row])], ignore_index=True)
        save_mood(mood_df)
        st.success("Eintrag gespeichert!")
        safe_rerun()

    st.markdown("---")
    st.subheader("📊 Verlauf der letzten 14 Tage")

    if not mood_df.empty:
        mood_df = mood_df.sort_values("datum")
        last_days = mood_df[mood_df["datum"] >= (today - timedelta(days=14))]

        if not last_days.empty:
            chart_data = last_days.set_index("datum")[["stimmung", "stress"]]
            st.line_chart(chart_data)

            st.markdown("### Letzte Einträge")
            st.dataframe(last_days.tail(20), use_container_width=True)

        st.markdown("---")
        st.subheader("🗑️ Falschen Eintrag löschen")

        mood_df_sorted = mood_df.sort_values("datum", ascending=False).reset_index().rename(columns={"index": "orig_index"})

        options = [
            f"{row['datum'].strftime('%d.%m.%Y')} – Stimmung: {row['stimmung']}/10, "
            f"Stress: {row['stress']}/10, Schlaf: {row['schlaf']}h"
            for _, row in mood_df_sorted.iterrows()
        ]

        selected_label = st.selectbox(
            "Eintrag auswählen, der gelöscht werden soll:",
            ["(kein Eintrag ausgewählt)"] + options,
        )

        if selected_label != "(kein Eintrag ausgewählt)":
            selected_idx = options.index(selected_label)
            row_to_delete = mood_df_sorted.iloc[selected_idx]

            st.warning(f"Du bist dabei, den Eintrag vom {row_to_delete['datum'].strftime('%d.%m.%Y')} zu löschen.")

            if st.button("❌ Ausgewählten Eintrag wirklich löschen"):
                orig_index = int(row_to_delete["orig_index"])
                mood_df = mood_df.drop(index=orig_index).reset_index(drop=True)
                save_mood(mood_df)
                st.success("Eintrag wurde gelöscht.")
                safe_rerun()

        st.markdown("---")
        st.subheader("🧠 Analyse & Hinweise")

        latest = mood_df.iloc[-1]
        l_stress = latest["stress"]
        l_schlaf = latest["schlaf"]
        l_stimmung = latest["stimmung"]

        if l_stress >= 8 and l_schlaf <= 5:
            st.error("Sehr hoher Stress und wenig Schlaf.\n\n👉 Versuche heute bewusst Pausen zu machen & früher zu schlafen.")
        elif l_stress >= 7:
            st.warning("Dein Stresslevel ist aktuell hoch.\n\n👉 Plane kleine Pausen ein.")
        elif l_stimmung <= 4:
            st.info("Deine Stimmung ist etwas im Keller.\n\n👉 Vielleicht hilft Bewegung/Musik/reden.")
        else:
            st.success("Alles im grünen Bereich – gute Voraussetzungen fürs Lernen! 💪")
    else:
        st.info("Noch keine Mood-Daten vorhanden. Mach oben deinen ersten Eintrag.")
